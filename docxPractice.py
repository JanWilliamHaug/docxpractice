#1. remember to install "python-docx" and "xlwings" first
#2. Import libraries
import xlwings as xw
import docx
from docx import Document
from tkinter import *
import re
# from docx2python import docx2python
import os
#import pandas as pd
#import mlx.traceability


excelFile = xw.Book("RTM.xlsx")
excelFile.save('RTM.xlsx')

#extensions = ['mlx.traceability']

#excelFile = xw.Book()                #Creates an empty excel file
#excelFile.save('report.xlsx')                     #Saves that excel file as "data1"

#ws1 = excelFile.sheets['Sheet1']

#Create a Docx file document object and pass the path to the Docx file
Text = docx.Document('SRS_ACE_Pump_X01.docx')

#staticPath = [os.path.join(os.path.dirname(mlx.traceability.Text), 'Software Requirement')]
'''traceAttributes = {
    'value': '^.*$',
    'asil': '^(QM|[ABCD])$',
    'non_functional': '^.{0}$',  # empty string
}

traceabiltiy_relationship_to_string = {
    'value': 'ACE:SRS:1',
    'asil': '[PUMP:PRS:1]',
}'''

#Create an empty data dictionary
#data = {}

#Create a paragraph object out of the document object.
#This object can access all the paragraphs of the document
#paragraphs = Text.paragraphs

# iterate over all the paragraphs, access the text, and save them into a data dictionary
#for i in range(0, len(Text.paragraphs)):
    #data[i] = tuple(Text.paragraphs[i].text.split(':'))

#values of the dictionary (list)
#data_values = list(data.values())
#print(data_values)

#m = re.search("SRS:(\w+)", paragraphs)
#print m.groups()
fullText = []
for para in Text.paragraphs:
    fullText.append(para.text)

mystring =' '.join(map(str, fullText))
print(mystring)

#pd.DataFrame(Text.body[1][1:])



product = "TARGEST"
name = "NAME"
coFounder1 = "Jan"
coFounder2 = "Adrian"
coFounder3 = "Stephania"
title = "TITLE"
title2 = "Co-Founder"

#check if there is a keyword that you are looking for and if it is, it will replace with the name
def find_(paragraph_keyword,paragraph):

    if paragraph_keyword in paragraph.text:
        print("found tag:", paragraph_keyword)
        #prints out "found tag:" whenever a tag is found

#going in the document.paragraphs using for loop
for paragraph in Text.paragraphs:

    find_("ACE", paragraph)
    find_("SRS", paragraph)
    find_("1", paragraph)
    find_("2", paragraph)
    find_("5", paragraph)
    find_("6", paragraph)
    find_("10", paragraph)
    find_("100", paragraph)
    find_("105", paragraph)
    find_("110", paragraph)
    find_("120", paragraph)
    find_("PUMP", paragraph)
    find_("PRS", paragraph)
    find_("TBV", paragraph)
    find_("DER", paragraph)
    find_("1000", paragraph)
    find_("Jan", paragraph)
    find_("CSU", paragraph)
    find_(':', paragraph)

    ws1 = excelFile.sheets['Sheet1']
    ws1.range('A3').value = product
    ws1.range('B3').value = coFounder1 + ", " + coFounder2 + "," + coFounder3
    ws1.range('A5').value = "SRS"
    ws1.range('A6').value = "SRS"
    ws1.range('A7').value = "SRS"
    ws1.range('A8').value = "SRS"
    ws1.range('A9').value = "SRS"
    ws1.range('A10').value = "SRS"
    ws1.range('A11').value = "SRS"
    ws1.range('A12').value = "SRS"
    ws1.range('A13').value = "PRS," + "TBV"
    ws1.range('A14').value = "PRS"
    ws1.range('A15').value = "PRS"
    ws1.range('A16').value = "PRS"
    ws1.range('A17').value = "PRS"
    ws1.range('A18').value = "PRS"
    ws1.range('A19').value = "PRS"
    ws1.range('A20').value = "DER"
    ws1.range('B5').value = "ACE" + ':' + "SRS" + ':' + "1"
    ws1.range('B6').value = "ACE" + ':' + "SRS" + ':' + "2"
    ws1.range('B7').value = "ACE" + ':' + "SRS" + ':' + "5"
    ws1.range('B8').value = "ACE" + ':' + "SRS" + ':' + "6"
    ws1.range('B9').value = "ACE" + ':' + "SRS" + ':' + "10"
    ws1.range('B10').value = "ACE" + ':' + "SRS" + ':' + "100"
    ws1.range('B11').value = "ACE" + ':' + "SRS" + ':' + "110"
    ws1.range('B12').value = "ACE" + ':' + "SRS" + ':' + "120"
    ws1.range('B13').value = "[PUMP" + ':' + "PRS" + ':' + "1]," + "[PUMP" + ':' + "TBV" + ':' + "1]"
    ws1.range('B14').value = "[PUMP" + ':' + "PRS" + ':' + "1]"
    ws1.range('B15').value = "[PUMP" + ':' + "PRS" + ':' + "5]"
    ws1.range('B16').value = "[PUMP" + ':' + "PRS" + ':' + "6]"
    ws1.range('B17').value = "[PUMP" + ':' + "PRS" + ':' + "10]"
    ws1.range('B18').value = "[PUMP" + ':' + "PRS" + ':' + "105]"
    ws1.range('B19').value = "[PUMP" + ':' + "PRS" + ':' + "1000]"
    ws1.range('B20').value = "[PUMP" + ':' + "DER" + ':' + "1]"


user = [re.findall('(?<=SRS: )\w+', s) for s in mystring]

print(user)
print('found tag', user)

